"""Word node — download → text extraction → Claude."""

from __future__ import annotations

import asyncio
import logging
from pathlib import Path

from core.downloader import download_to_temp
from core.llm import extract_from_content
from core.timing import FileTimer
from models import ExtractionError, ExtractedContent, FileItem, FileType
from state import CIMState

log = logging.getLogger(__name__)


def _docx_to_text(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            lines.append("\t".join(cell.text.strip() for cell in row.cells))
    return "\n".join(lines)


def _rtf_to_text(path: Path) -> str:
    import re
    raw = path.read_bytes().decode("utf-8", errors="ignore")
    text = re.sub(r"\{\\[^{}]+\}", "", raw)
    text = re.sub(r"\\[a-z]+\d*\s?", "", text)
    text = re.sub(r"[{}]", "", text)
    return text.strip()


def _to_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in (".docx", ".docm"):
        return _docx_to_text(path)
    if suffix == ".doc":
        try:
            return _docx_to_text(path)
        except Exception:
            pass
        return path.read_text(errors="ignore")
    if suffix == ".odt":
        try:
            return _docx_to_text(path)
        except Exception:
            return path.read_text(errors="ignore")
    if suffix == ".rtf":
        return _rtf_to_text(path)
    return path.read_text(errors="ignore")


async def _process_one(item: FileItem, listing_xml: str = "") -> ExtractedContent:
    label = item.label or item.url.split("/")[-1].split("?")[0]
    try:
        with FileTimer("word", label):
            if item.content:
                import tempfile, base64 as _b64
                suffix = "." + item.label.rsplit(".", 1)[-1] if "." in item.label else ".docx"
                with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                    tmp.write(_b64.b64decode(item.content))
                    tmp_path = Path(tmp.name)
                try:
                    raw_text = await asyncio.to_thread(_to_text, tmp_path)
                finally:
                    tmp_path.unlink(missing_ok=True)
            else:
                async with download_to_temp(item.url) as path:
                    raw_text = await asyncio.to_thread(_to_text, path)

            content_blocks = [{"type": "text", "text": raw_text}]
            text = await extract_from_content(content_blocks, item.url, listing_xml)
        log.debug("  Word extracted: %r", label)
        return ExtractedContent(
            source_url=item.url,
            file_type=FileType.WORD,
            metadata={"extracted_text": text},
        )
    except Exception as exc:
        log.error("  Word failed: %r — %s", label, exc)
        return ExtractedContent(source_url=item.url, file_type=FileType.WORD, error=str(exc))


async def word_node(state: CIMState) -> dict:
    files: list[FileItem] = state.get("word_files", [])
    if not files:
        return {"extracted": [], "errors": []}

    log.debug("Word: processing %d file(s)", len(files))
    results = await asyncio.gather(*[_process_one(f, state.get("listing_xml", "")) for f in files])

    extracted = [r for r in results if not r.error]
    errors = [ExtractionError(url=r.source_url, stage="word", message=r.error) for r in results if r.error]
    if errors:
        log.error("Word: %d error(s)", len(errors))
    return {"extracted": list(extracted), "errors": errors}
